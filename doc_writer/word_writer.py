from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from .base_writer import BaseDocumentWriter

_A4_WIDTH_CM = 21.0
_A4_HEIGHT_CM = 29.7

_ALIGN_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


# ---------------------------------------------------------------------------
# Low-level helpers
# ---------------------------------------------------------------------------

def _rgb(hex_color: str) -> RGBColor:
    h = hex_color.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _apply_font(run, font_def: dict) -> None:
    run.font.name = font_def.get("name", "Calibri")
    run.font.size = Pt(font_def.get("size", 11))
    run.font.bold = font_def.get("bold", False)
    run.font.italic = font_def.get("italic", False)
    color = font_def.get("color")
    if color:
        run.font.color.rgb = _rgb(color)


def _is_cell(container) -> bool:
    return hasattr(container, "_tc")


def _para_in(container):
    """Return the first paragraph of a cell if unused, else add a new one."""
    if _is_cell(container):
        paras = container.paragraphs
        if paras and paras[0].text == "":
            return paras[0]
    return container.add_paragraph()


def _add_table(container, doc: Document, rows: int, cols: int):
    if _is_cell(container):
        return container.add_table(rows, cols)
    return doc.add_table(rows=rows, cols=cols)


def _remove_table_borders(table) -> None:
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        borders.append(el)
    tblPr.append(borders)


def _set_cell_width(cell, twips: int) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:tcW")):
        tcPr.remove(old)
    w = OxmlElement("w:tcW")
    w.set(qn("w:w"), str(twips))
    w.set(qn("w:type"), "dxa")
    tcPr.append(w)


def _set_cell_shading(cell, fill_hex: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:shd")):
        tcPr.remove(old)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex.lstrip("#"))
    tcPr.append(shd)


def _set_cell_borders(cell, color_hex: str, width_pt: float) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:tcBorders")):
        tcPr.remove(old)
    tcBorders = OxmlElement("w:tcBorders")
    sz = str(int(width_pt * 8))
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color_hex.lstrip("#"))
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _set_callout_borders(cell, color_hex: str, width_pt: float) -> None:
    """Left border only — used for callout blocks."""
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:tcBorders")):
        tcPr.remove(old)
    tcBorders = OxmlElement("w:tcBorders")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(int(width_pt * 8)))
    left.set(qn("w:space"), "0")
    left.set(qn("w:color"), color_hex.lstrip("#"))
    tcBorders.append(left)
    for side in ("top", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        tcBorders.append(el)
    tcPr.append(tcBorders)


# ---------------------------------------------------------------------------
# Writer
# ---------------------------------------------------------------------------

class WordDocumentWriter(BaseDocumentWriter):
    """Generates a .docx document from doc_writer templates."""

    def __init__(self, custom_templates_dir: Path | str | None = None):
        super().__init__(custom_templates_dir)
        self._doc = Document()
        self._content_width_cm: float = _A4_WIDTH_CM - 6.0  # default margins 3cm each side

    # ------------------------------------------------------------------
    # Page setup
    # ------------------------------------------------------------------

    def _setup_page(self, page_def: dict) -> None:
        mm = page_def.get("margins_mm", {})
        top = mm.get("top", 25) / 10
        bottom = mm.get("bottom", 25) / 10
        left = mm.get("left", 30) / 10
        right = mm.get("right", 30) / 10
        sec = self._doc.sections[0]
        sec.page_width = Cm(_A4_WIDTH_CM)
        sec.page_height = Cm(_A4_HEIGHT_CM)
        sec.top_margin = Cm(top)
        sec.bottom_margin = Cm(bottom)
        sec.left_margin = Cm(left)
        sec.right_margin = Cm(right)
        self._content_width_cm = _A4_WIDTH_CM - left - right

    # ------------------------------------------------------------------
    # Layout rendering
    # ------------------------------------------------------------------

    def _render_combined(self, resolved_rows: list[dict]) -> None:
        for row in resolved_rows:
            cols = row["columns"]
            if len(cols) == 1:
                for comp in cols[0]["components"]:
                    self._render_component(self._doc, comp)
            else:
                # Invisible layout table for multi-column rows
                table = self._doc.add_table(rows=1, cols=len(cols))
                _remove_table_borders(table)
                total_twips = int(Cm(self._content_width_cm).pt * 20)
                for idx, col in enumerate(cols):
                    cell = table.rows[0].cells[idx]
                    _set_cell_width(cell, int(total_twips * col["width_pct"] / 100))
                    for comp in col["components"]:
                        self._render_component(cell, comp)

    def _render_component(self, container, comp: dict) -> None:
        tpl = comp["template"]
        content = comp["content"]
        dispatch = {
            "title": self._write_title,
            "paragraph": self._write_paragraph,
            "bullet_list": self._write_bullet_list,
            "image": self._write_image,
            "table": self._write_table,
            "callout": self._write_callout,
            "separator": self._write_separator,
            "chart": self._write_chart,
            "flowchart": self._write_flowchart,
        }
        handler = dispatch.get(tpl["type"])
        if handler:
            handler(container, tpl, content)

    # ------------------------------------------------------------------
    # Singular renderers
    # ------------------------------------------------------------------

    def _write_title(self, container, tpl: dict, content: dict) -> None:
        text = content.get("text", "")
        level = str(content.get("level", 1))
        level_def = tpl["levels"].get(level, next(iter(tpl["levels"].values())))
        font_def = level_def["font"]
        sp = level_def.get("spacing", {})

        para = _para_in(container)
        para.alignment = _ALIGN_MAP.get(font_def.get("alignment", "left"), WD_ALIGN_PARAGRAPH.LEFT)
        para.paragraph_format.space_before = Pt(sp.get("before_pt", 12))
        para.paragraph_format.space_after = Pt(sp.get("after_pt", 6))
        _apply_font(para.add_run(text), font_def)

    def _write_paragraph(self, container, tpl: dict, content: dict) -> None:
        text = content.get("text", "")
        font_def = tpl["font"]
        sp = tpl.get("spacing", {})

        para = _para_in(container)
        para.alignment = _ALIGN_MAP.get(tpl.get("alignment", "left"), WD_ALIGN_PARAGRAPH.LEFT)
        para.paragraph_format.space_before = Pt(sp.get("before_pt", 6))
        para.paragraph_format.space_after = Pt(sp.get("after_pt", 6))
        _apply_font(para.add_run(text), font_def)

    def _write_bullet_list(self, container, tpl: dict, content: dict) -> None:
        items = content.get("items", [])
        bullet_def = tpl["bullet"]
        font_def = tpl["font"]
        sp = tpl.get("spacing", {})
        indent = tpl.get("indent_pt", 18)

        for item in items:
            para = container.add_paragraph()
            para.alignment = _ALIGN_MAP.get(tpl.get("alignment", "left"), WD_ALIGN_PARAGRAPH.LEFT)
            para.paragraph_format.left_indent = Pt(indent)
            para.paragraph_format.first_line_indent = Pt(-indent)
            para.paragraph_format.space_before = Pt(sp.get("before_pt", 2))
            para.paragraph_format.space_after = Pt(sp.get("after_pt", 2))

            bullet_run = para.add_run(f"{bullet_def['symbol']}  ")
            bullet_run.font.size = Pt(bullet_def.get("size", 11))
            if c := bullet_def.get("color"):
                bullet_run.font.color.rgb = _rgb(c)

            _apply_font(para.add_run(str(item)), font_def)

    def _write_image(self, container, tpl: dict, content: dict) -> None:
        frame = tpl.get("frame", {})
        caption_def = tpl.get("caption", {})
        path = content.get("path")
        caption_text = content.get("caption", "")

        if path:
            from docx.shared import Inches
            para = _para_in(container)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.add_run().add_picture(str(path), width=Inches(4))
        else:
            placeholder = tpl.get("placeholder_label", "[ IMAGE ]")
            img_tbl = _add_table(container, self._doc, 1, 1)
            ph_cell = img_tbl.rows[0].cells[0]
            _set_cell_borders(ph_cell, frame.get("border_color", "#CCCCCC"), frame.get("border_width_pt", 1.5))
            _set_cell_shading(ph_cell, frame.get("background_color", "#FAFAFA"))
            p = ph_cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"\n{placeholder}\n")
            run.font.color.rgb = RGBColor(153, 153, 153)
            run.font.italic = True
            run.font.size = Pt(12)

        if caption_text:
            cap_para = container.add_paragraph()
            cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _apply_font(cap_para.add_run(caption_text), caption_def.get("font", {}))

    def _write_table(self, container, tpl: dict, content: dict) -> None:
        headers = content.get("headers", [])
        rows_data = content.get("rows", [])
        header_def = tpl.get("header", {})
        cell_def = tpl.get("cell", {})
        border_def = tpl.get("border", {})

        num_cols = max(
            tpl.get("columns", 3),
            len(headers),
            max((len(r) for r in rows_data), default=0),
        )
        num_rows = (1 if headers else 0) + len(rows_data)
        if num_rows == 0 or num_cols == 0:
            return

        table = _add_table(container, self._doc, num_rows, num_cols)
        border_color = border_def.get("color", "#AAAAAA")
        border_width = border_def.get("width_pt", 0.5)

        row_offset = 0
        if headers:
            for col_idx, hdr in enumerate(headers[:num_cols]):
                cell = table.rows[0].cells[col_idx]
                _set_cell_borders(cell, border_color, border_width)
                _set_cell_shading(cell, header_def.get("background_color", "#333333"))
                p = cell.paragraphs[0]
                p.alignment = _ALIGN_MAP.get(header_def.get("alignment", "center"), WD_ALIGN_PARAGRAPH.CENTER)
                _apply_font(p.add_run(str(hdr)), header_def.get("font", {}))
            row_offset = 1

        for r_idx, row_data in enumerate(rows_data):
            for col_idx, val in enumerate(row_data[:num_cols]):
                cell = table.rows[row_offset + r_idx].cells[col_idx]
                _set_cell_borders(cell, border_color, border_width)
                padding = cell_def.get("padding_pt", 4)
                p = cell.paragraphs[0]
                p.alignment = _ALIGN_MAP.get(cell_def.get("alignment", "left"), WD_ALIGN_PARAGRAPH.LEFT)
                p.paragraph_format.space_before = Pt(padding)
                p.paragraph_format.space_after = Pt(padding)
                _apply_font(p.add_run(str(val)), cell_def.get("font", {}))

    def _write_callout(self, container, tpl: dict, content: dict) -> None:
        label = content.get("label", "")
        text = content.get("text", "")
        border_color = tpl.get("border_left_color", "#2980B9")
        bg_color = tpl.get("background_color", "#EBF5FB")
        border_width = tpl.get("border_left_width_pt", 4)
        sp = tpl.get("spacing", {})

        tbl = _add_table(container, self._doc, 1, 1)
        _remove_table_borders(tbl)
        cell = tbl.rows[0].cells[0]
        _set_callout_borders(cell, border_color, border_width)
        _set_cell_shading(cell, bg_color)

        if label:
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(sp.get("before_pt", 8))
            _apply_font(p.add_run(label), tpl.get("label_font", {}))
            p2 = cell.add_paragraph()
            p2.paragraph_format.space_after = Pt(sp.get("after_pt", 8))
            _apply_font(p2.add_run(text), tpl.get("font", {}))
        else:
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(sp.get("before_pt", 8))
            p.paragraph_format.space_after = Pt(sp.get("after_pt", 8))
            _apply_font(p.add_run(text), tpl.get("font", {}))

    def _write_separator(self, container, tpl: dict, content: dict) -> None:
        color = tpl.get("line_color", "#CCCCCC").lstrip("#")
        width_pt = tpl.get("line_width_pt", 0.5)
        margin_top = tpl.get("margin_top_pt", 12)
        margin_bottom = tpl.get("margin_bottom_pt", 12)

        para = container.add_paragraph()
        para.paragraph_format.space_before = Pt(margin_top)
        para.paragraph_format.space_after = Pt(margin_bottom)

        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), str(int(width_pt * 8)))
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), color)
        pBdr.append(bottom)
        pPr.append(pBdr)

    def _write_chart(self, container, tpl: dict, content: dict) -> None:
        from docx.shared import Inches
        from .chart_utils import render_chart_png
        buf = render_chart_png(tpl, content)
        para = container.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Pt(4)
        para.paragraph_format.space_after = Pt(4)
        width_in = tpl.get("figure", {}).get("width_inches", 6)
        para.add_run().add_picture(buf, width=Inches(width_in))

    def _write_flowchart(self, container, tpl: dict, content: dict) -> None:
        from docx.shared import Inches
        from .flowchart_utils import (build_simple_png_for_word,
                                      build_complex_png_for_word)
        ftype = content.get("flowchart_type", "simple")
        nodes = content.get("nodes", [])
        if ftype == "simple":
            buf = build_simple_png_for_word(nodes, tpl)
        else:
            buf = build_complex_png_for_word(nodes, content.get("edges", []), tpl)
        width_in = tpl.get("figure", {}).get("width_inches", 4.5)
        para = container.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Pt(4)
        para.paragraph_format.space_after = Pt(4)
        para.add_run().add_picture(buf, width=Inches(width_in))

    # ------------------------------------------------------------------
    # Save
    # ------------------------------------------------------------------

    def save(self, output_path: str | Path) -> None:
        out = Path(output_path)
        out.parent.mkdir(parents=True, exist_ok=True)
        self._doc.save(str(out))
